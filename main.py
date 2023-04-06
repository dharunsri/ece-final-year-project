from flask import Flask, request, render_template, send_file
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import cv2
import numpy as np
from keras.models import load_model
#import nltk
#nltk.download('punkt')
from nltk.tokenize import word_tokenize
import spacy
import os
from spellchecker import SpellChecker
from OCR import Prediction

app = Flask(__name__)

# Load the trained CNN model
model = load_model('cnn_model.h5')

# Load a more accurate NLP model for word replacement
nlp_trf = spacy.load("en_core_web_trf")
nlp_md = spacy.load('en_core_web_md')

# Define a function to preprocess the input image
def preprocess_image(image):
    # Convert the image to grayscale
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # Apply Gaussian blur to reduce noise
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    # Apply adaptive thresholding to binarize the image
    thresh = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 2)
    # Resize the image to match the input shape of the CNN model
    resized = cv2.resize(thresh, (28, 28), interpolation=cv2.INTER_AREA)
    # Reshape the image to a 4D tensor with a single channel
    tensor = np.expand_dims(np.expand_dims(resized, axis=-1), axis=0)
    # Normalize the tensor to have values between 0 and 1
    tensor = tensor.astype('float32') / 255.0
    return tensor

# Define the custom extensions
def has_suggestion(token):
    return token._.suggestions != []

def suggestion(token):
    if token._.has_extension('suggestion'):
        return token._.suggestion
    else:
        return None

def get_suggestion(token):
    if token._.has_extension('suggestion'):
        return token._.suggestion
    else:
        return None

# Register the custom extensions
spacy.tokens.Token.set_extension('suggestion', default=None)
spacy.tokens.Token.set_extension('get_suggestion', getter=get_suggestion)
spacy.tokens.Token.set_extension('has_suggestion', getter=has_suggestion)

def replace_misspelled_words(text):
    spell = SpellChecker()
    corrected_text = []
    misspelled_words = spell.unknown(text.split())

    for word in text.split():
        if word in misspelled_words:
            suggestion = spell.correction(word)
            if suggestion is not None:
                corrected_text.append(suggestion)
            else:
                corrected_text.append(word)
        else:
            corrected_text.append(word)

    return ' '.join(corrected_text)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Get the uploaded file and load it as a document object
        file = request.files['file']
        print(file)
        out_doc = Document()
        try:
           text = Prediction.recognise(file)
           text = replace_misspelled_words(text)
           out_para = out_doc.add_paragraph()

           for word in text.split(' '):
               out_para.add_run(word + ' ')
        
        except:
            doc = Document(file)
            # Loop through all paragraphs in the input document
            for para in doc.paragraphs:
                # Get the text content of the paragraph
                text = para.text
                # Replace misspelled words with correctly predicted words using NLP
                text = replace_misspelled_words(text)  # Call the function here
                # Add the modified text to the output document
                out_para = out_doc.add_paragraph()
                # Loop through all words in the modified text
                for word in text.split(' '):
                    # Check if the word was replaced (i.e., contains at least one non-letter character)
                    if not word.isalpha():
                        # Add the word to the output document with a highlight color
                        run = out_para.add_run(word)
                        #run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    else:
                        # Add the word to the output document
                        out_para.add_run(word + ' ')
        # Save the output document to a new file
        out_doc.save('output.docx')
        # Download the output file to the user's computer
        return send_file('output.docx', as_attachment=True)
    else:
        # Render the upload page for the user to select a file
        return render_template('index.html')
    
@app.route('/my_team')
def my_team():
    return render_template('team.html')

@app.route('/guide')
def guide():
    return render_template('guide.html')

@app.route('/project')
def project():
    return render_template('project.html')

if __name__ == '__main__':
    app.run(debug=True)
